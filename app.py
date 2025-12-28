import sys
import json
import requests
from datetime import datetime
from docx import Document
from PyQt6.QtWidgets import QFileDialog
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTabWidget, QMessageBox,
    QTableWidget, QTableWidgetItem, QHeaderView, QDialog, QFormLayout,
    QComboBox, QDateEdit, QAbstractItemView, QMenu
)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QAction, QIcon, QFont

# Настройки API
API_BASE_URL = "http://localhost:5000"
HEADERS = {}


class LoginDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Вход в систему")
        self.setFixedSize(400, 200)
        self.parent = parent

        layout = QVBoxLayout()
        login_layout = QHBoxLayout()
        login_layout.addWidget(QLabel("Логин:"))
        self.login_input = QLineEdit()
        login_layout.addWidget(self.login_input)
        layout.addLayout(login_layout)

        pass_layout = QHBoxLayout()
        pass_layout.addWidget(QLabel("Пароль:"))
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        pass_layout.addWidget(self.password_input)
        layout.addLayout(pass_layout)

        self.login_button = QPushButton("Войти")
        self.login_button.clicked.connect(self.attempt_login)
        layout.addWidget(self.login_button)
        self.setLayout(layout)

    def attempt_login(self):
        login = self.login_input.text().strip()
        password = self.password_input.text().strip()
        if not login or not password:
            QMessageBox.critical(self, "Ошибка", "Введите логин и пароль!")
            return

        try:
            response = requests.post(
                f"{API_BASE_URL}/login",
                json={"login": login, "password": password},
                timeout=10
            )
            if response.status_code == 200:
                data = response.json()
                self.parent.current_user_id = data["user_id"]
                self.parent.current_user_role = data["role"]
                HEADERS['X-User-ID'] = str(data["user_id"])
                self.accept()
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                QMessageBox.critical(self, "Ошибка", f"Неверный логин или пароль.\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось подключиться к серверу:\n{str(e)}")


class EntityTableWidget(QTableWidget):
    def __init__(self, parent, entity_type, columns, main_window, editable_fields=None):
        super().__init__(parent)
        self.entity_type = entity_type
        self.columns = columns
        self.main_window = main_window
        self.editable_fields = editable_fields or []
        self.setColumnCount(len(columns))
        self.setHorizontalHeaderLabels(columns)
        self.horizontalHeader().setStretchLastSection(True)
        self.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.cellDoubleClicked.connect(self.on_cell_double_clicked)

    def on_cell_double_clicked(self, row, col):
        # Арендатор не может редактировать ничего
        if self.main_window.current_user_role == "tenant":
            return
        if self.main_window.current_user_role == "accountant" and self.entity_type == "building":
            return

        item = self.item(row, 0)
        if item:
            entity_id = int(item.text())
            self.main_window.edit_entity(self.entity_type, entity_id)

    def populate_table(self, data_list):
        self.setRowCount(0)
        for data in data_list:
            if not isinstance(data, dict):
                continue
            row_position = self.rowCount()
            self.insertRow(row_position)
            for col_idx, field in enumerate(self.columns):
                value = ""
                if field == "ID":
                    value = str(data.get("id", ""))
                elif field == "Название":
                    value = data.get("name", "")
                elif field == "Часовой пояс":
                    value = data.get("timezone", "")
                elif field == "Цена за кВт·ч":
                    value = str(data.get("rate_per_kwh", ""))
                elif field == "С":
                    value = data.get("valid_from", "")
                elif field == "По":
                    value = data.get("valid_to", "") or "-"
                elif field == "Логин":
                    value = data.get("login", "")
                elif field == "Роль":
                    value = data.get("role", "")
                elif field == "Адрес":
                    value = data.get("address", "")
                elif field == "Тип":
                    value = data.get("type", "")
                elif field == "Регион":
                    value = data.get("region_name", "") or ""
                elif field == "Тариф":
                    value = data.get("tariff_name", "") or ""
                elif field == "Владелец":
                    value = data.get("owner_login", "") or ""
                elif field == "Серийный номер":
                    value = data.get("serial_number", "") or ""
                elif field == "Дата установки":
                    value = data.get("installation_date", "") or ""
                elif field == "Период с":
                    value = data.get("period_start", "") or ""
                elif field == "Период по":
                    value = data.get("period_end", "") or ""
                elif field == "кВт·ч":
                    value = str(data.get("consumption_kwh", ""))
                elif field == "Оценка (руб)":
                    cost = data.get("estimated_cost_rub")
                    value = str(cost) if cost is not None else "-"
                elif field == "Счётчик":
                    meter_serial = data.get("meter_serial")
                    value = meter_serial if meter_serial else str(data.get("meter_id", ""))
                elif field == "Объект":
                    # Для вкладки "Счётчики"
                    value = data.get("building_name") \
                            or data.get("building") \
                            or str(data.get("building_id", ""))

                item = QTableWidgetItem(value)
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                self.setItem(row_position, col_idx, item)


class EditEntityDialog(QDialog):
    def __init__(self, parent, entity_type, entity_data, main_window):
        super().__init__(parent)
        self.setWindowTitle(f"Редактировать {entity_type}")
        self.entity_type = entity_type
        self.entity_data = entity_data
        self.parent = parent
        self.main_window = main_window

        layout = QFormLayout()
        self.id_field = QLineEdit(str(entity_data["id"]))
        self.id_field.setVisible(False)
        layout.addRow("", self.id_field)

        if entity_type == "region":
            self.name_field = QLineEdit(entity_data.get("name", ""))
            self.timezone_field = QLineEdit(entity_data.get("timezone", ""))
            layout.addRow("Название:", self.name_field)
            layout.addRow("Часовой пояс:", self.timezone_field)
        elif entity_type == "tariff":
            self.name_field = QLineEdit(entity_data.get("name", ""))
            self.rate_field = QLineEdit(str(entity_data.get("rate_per_kwh", "")))
            self.valid_from_field = QDateEdit()
            self.valid_from_field.setDate(
                datetime.fromisoformat(entity_data.get("valid_from")).date()
                if entity_data.get("valid_from") else datetime.now().date()
            )
            self.valid_to_field = QDateEdit()
            valid_to = entity_data.get("valid_to")
            if valid_to:
                self.valid_to_field.setDate(datetime.fromisoformat(valid_to).date())
            else:
                self.valid_to_field.setDate(datetime.now().date())
            layout.addRow("Название:", self.name_field)
            layout.addRow("Цена за кВт·ч:", self.rate_field)
            layout.addRow("Действует с:", self.valid_from_field)
            layout.addRow("Действует до:", self.valid_to_field)
        elif entity_type == "building":
            self.name_field = QLineEdit(entity_data.get("name", ""))
            self.address_field = QLineEdit(entity_data.get("address", ""))
            self.type_combo = QComboBox()
            self.type_combo.addItems(["жилое", "промышленное", "общественное"])
            self.type_combo.setCurrentText(entity_data.get("type", "жилое"))
            self.region_combo = QComboBox()
            self.tariff_combo = QComboBox()
            self.user_combo = QComboBox()
            self.load_dropdowns()
            layout.addRow("Название:", self.name_field)
            layout.addRow("Адрес:", self.address_field)
            layout.addRow("Тип:", self.type_combo)
            layout.addRow("Регион:", self.region_combo)
            layout.addRow("Тариф:", self.tariff_combo)
            layout.addRow("Владелец:", self.user_combo)
        elif entity_type == "meter":
            self.serial_field = QLineEdit(entity_data.get("serial_number", ""))
            self.installation_date_field = QDateEdit()
            self.installation_date_field.setDate(
                datetime.fromisoformat(entity_data.get("installation_date")).date()
                if entity_data.get("installation_date") else datetime.now().date()
            )
            self.building_combo = QComboBox()
            self.load_buildings()
            layout.addRow("Серийный номер:", self.serial_field)
            layout.addRow("Дата установки:", self.installation_date_field)
            layout.addRow("Объект учёта:", self.building_combo)
        elif entity_type == "consumption":
            self.period_start_field = QDateEdit()
            self.period_start_field.setDate(
                datetime.fromisoformat(entity_data.get("period_start")).date()
                if entity_data.get("period_start") else datetime.now().date()
            )
            self.period_end_field = QDateEdit()
            self.period_end_field.setDate(
                datetime.fromisoformat(entity_data.get("period_end")).date()
                if entity_data.get("period_end") else datetime.now().date()
            )
            self.kwh_field = QLineEdit(str(entity_data.get("consumption_kwh", "")))
            self.meter_combo = QComboBox()
            self.load_meters()
            layout.addRow("Счётчик:", self.meter_combo)
            layout.addRow("Период с:", self.period_start_field)
            layout.addRow("Период по:", self.period_end_field)
            layout.addRow("кВт·ч:", self.kwh_field)
        elif entity_type == "user":
            self.login_field = QLineEdit(entity_data.get("login", ""))
            self.role_combo = QComboBox()
            self.role_combo.addItems(["tenant", "accountant", "admin"])
            current_role = entity_data.get("role", "tenant")
            idx = self.role_combo.findText(current_role)
            if idx >= 0:
                self.role_combo.setCurrentIndex(idx)
            layout.addRow("Логин:", self.login_field)
            layout.addRow("Роль:", self.role_combo)

        button_layout = QHBoxLayout()
        save_btn = QPushButton("Сохранить")
        save_btn.clicked.connect(self.save_changes)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(save_btn)
        button_layout.addWidget(cancel_btn)
        layout.addRow(button_layout)
        self.setLayout(layout)

    def load_dropdowns(self):
        try:
            regions = requests.get(f"{API_BASE_URL}/regions", headers=HEADERS).json()
            tariffs = requests.get(f"{API_BASE_URL}/tariffs", headers=HEADERS).json()
            users = requests.get(f"{API_BASE_URL}/users", headers=HEADERS).json()
            if not isinstance(regions, list) or not isinstance(tariffs, list) or not isinstance(users, list):
                raise ValueError("Invalid data format from API")
            self.region_combo.clear()
            for r in regions:
                self.region_combo.addItem(r["name"], r["id"])
            self.tariff_combo.clear()
            for t in tariffs:
                self.tariff_combo.addItem(t["name"], t["id"])
            self.user_combo.clear()
            for u in users:
                self.user_combo.addItem(f"{u['login']} ({u['role']})", u["id"])

            current_region_id = self.entity_data.get("region_id")
            idx = self.region_combo.findData(current_region_id)
            if idx >= 0:
                self.region_combo.setCurrentIndex(idx)

            current_tariff_id = self.entity_data.get("tariff_id")
            idx = self.tariff_combo.findData(current_tariff_id)
            if idx >= 0:
                self.tariff_combo.setCurrentIndex(idx)

            current_user_id = self.entity_data.get("user_id")
            idx = self.user_combo.findData(current_user_id)
            if idx >= 0:
                self.user_combo.setCurrentIndex(idx)
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные:\n{str(e)}")

    def load_buildings(self):
        try:
            if self.main_window.current_user_role == "tenant":
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(buildings, list):
                    raise ValueError("Invalid data format for buildings")
                user_id = self.main_window.current_user_id
                buildings = [b for b in buildings if b.get("user_id") == user_id]
            else:
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(buildings, list):
                    raise ValueError("Invalid data format for buildings")

            self.building_combo.clear()
            for b in buildings:
                self.building_combo.addItem(b["name"], b["id"])
            current_building_id = self.entity_data.get("building_id")
            idx = self.building_combo.findData(current_building_id)
            if idx >= 0:
                self.building_combo.setCurrentIndex(idx)
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить объекты:\n{str(e)}")

    def load_meters(self):
        try:
            if self.main_window.current_user_role == "tenant":
                meters = requests.get(f"{API_BASE_URL}/meters", headers=HEADERS).json()
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(meters, list) or not isinstance(buildings, list):
                    raise ValueError("Invalid data format")
                user_id = self.main_window.current_user_id
                building_ids = [b["id"] for b in buildings if b.get("user_id") == user_id]
                meters = [m for m in meters if m.get("building_id") in building_ids]
            else:
                meters = requests.get(f"{API_BASE_URL}/meters", headers=HEADERS).json()
                if not isinstance(meters, list):
                    raise ValueError("Invalid data format for meters")

            self.meter_combo.clear()
            for m in meters:
                self.meter_combo.addItem(m["serial_number"], m["id"])
            current_meter_id = self.entity_data.get("meter_id")
            idx = self.meter_combo.findData(current_meter_id)
            if idx >= 0:
                self.meter_combo.setCurrentIndex(idx)
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить счётчики:\n{str(e)}")

    def save_changes(self):
        data = {}
        if self.entity_type == "region":
            data = {
                "name": self.name_field.text().strip(),
                "timezone": self.timezone_field.text().strip()
            }
        elif self.entity_type == "tariff":
            try:
                rate = float(self.rate_field.text().strip())
            except ValueError:
                QMessageBox.critical(self, "Ошибка", "Цена за кВт·ч должна быть числом!")
                return
            data = {
                "name": self.name_field.text().strip(),
                "rate_per_kwh": rate,
                "valid_from": self.valid_from_field.date().toString("yyyy-MM-dd"),
                "valid_to": self.valid_to_field.date().toString("yyyy-MM-dd") if self.valid_to_field.date().isValid() else None
            }
        elif self.entity_type == "building":
            region_id = self.region_combo.currentData()
            tariff_id = self.tariff_combo.currentData()
            user_id = self.user_combo.currentData()
            if not region_id or not tariff_id or not user_id:
                QMessageBox.critical(self, "Ошибка", "Выберите регион, тариф и владельца!")
                return
            data = {
                "name": self.name_field.text().strip(),
                "address": self.address_field.text().strip(),
                "type": self.type_combo.currentText(),
                "region_id": region_id,
                "tariff_id": tariff_id,
                "user_id": user_id
            }
        elif self.entity_type == "meter":
            building_id = self.building_combo.currentData()
            if not building_id:
                QMessageBox.critical(self, "Ошибка", "Выберите объект учёта!")
                return
            data = {
                "serial_number": self.serial_field.text().strip(),
                "installation_date": self.installation_date_field.date().toString("yyyy-MM-dd"),
                "building_id": building_id
            }
        elif self.entity_type == "consumption":
            meter_id = self.meter_combo.currentData()
            if not meter_id:
                QMessageBox.critical(self, "Ошибка", "Выберите счётчик!")
                return
            try:
                kwh = float(self.kwh_field.text().strip())
            except ValueError:
                QMessageBox.critical(self, "Ошибка", "Потребление должно быть числом!")
                return
            data = {
                "meter_id": meter_id,
                "period_start": self.period_start_field.date().toString("yyyy-MM-dd"),
                "period_end": self.period_end_field.date().toString("yyyy-MM-dd"),
                "consumption_kwh": kwh
            }
        elif self.entity_type == "user":
            login = self.login_field.text().strip()
            role = self.role_combo.currentText()
            if not login:
                QMessageBox.critical(self, "Ошибка", "Введите логин!")
                return
            data = {
                "login": login,
                "role_id": 1 if role == "tenant" else (2 if role == "accountant" else 3)
            }

        try:
            if self.entity_type == "consumption":
                url = f"{API_BASE_URL}/consumption/{self.entity_data['id']}"
            else:
                url = f"{API_BASE_URL}/{self.entity_type}s/{self.entity_data['id']}"

            response = requests.put(url, json=data, headers=HEADERS, timeout=10)
            if response.status_code == 200:
                QMessageBox.information(self, "Успех", "Запись успешно обновлена!")
                self.accept()
                self.parent.refresh_data(self.entity_type)
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить:\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка подключения:\n{str(e)}")


class AddEntityDialog(QDialog):
    def __init__(self, parent, entity_type, main_window):
        super().__init__(parent)
        self.setWindowTitle(f"Добавить {entity_type}")
        self.entity_type = entity_type
        self.parent = parent
        self.main_window = main_window
        layout = QFormLayout()

        if entity_type == "region":
            self.name_field = QLineEdit()
            self.timezone_field = QLineEdit()
            layout.addRow("Название:", self.name_field)
            layout.addRow("Часовой пояс:", self.timezone_field)
        elif entity_type == "tariff":
            self.name_field = QLineEdit()
            self.rate_field = QLineEdit()
            self.valid_from_field = QDateEdit()
            self.valid_from_field.setDate(datetime.now().date())
            self.valid_to_field = QDateEdit()
            self.valid_to_field.setDate(datetime.now().date())
            layout.addRow("Название:", self.name_field)
            layout.addRow("Цена за кВт·ч:", self.rate_field)
            layout.addRow("Действует с:", self.valid_from_field)
            layout.addRow("Действует до:", self.valid_to_field)
        elif entity_type == "user":
            self.login_field = QLineEdit()
            self.role_combo = QComboBox()
            self.role_combo.addItems(["tenant", "accountant", "admin"])
            layout.addRow("Логин:", self.login_field)
            layout.addRow("Роль:", self.role_combo)
        elif entity_type == "building":
            self.name_field = QLineEdit()
            self.address_field = QLineEdit()
            self.type_combo = QComboBox()
            self.type_combo.addItems(["жилое", "промышленное", "общественное"])
            self.region_combo = QComboBox()
            self.tariff_combo = QComboBox()
            self.user_combo = QComboBox()
            self.load_dropdowns()
            layout.addRow("Название:", self.name_field)
            layout.addRow("Адрес:", self.address_field)
            layout.addRow("Тип:", self.type_combo)
            layout.addRow("Регион:", self.region_combo)
            layout.addRow("Тариф:", self.tariff_combo)
            layout.addRow("Владелец:", self.user_combo)
        elif entity_type == "meter":
            self.serial_field = QLineEdit()
            self.installation_date_field = QDateEdit()
            self.installation_date_field.setDate(datetime.now().date())
            self.building_combo = QComboBox()
            self.load_buildings()
            layout.addRow("Серийный номер:", self.serial_field)
            layout.addRow("Дата установки:", self.installation_date_field)
            layout.addRow("Объект учёта:", self.building_combo)
        elif entity_type == "consumption":
            self.meter_combo = QComboBox()
            self.period_start_field = QDateEdit()
            self.period_start_field.setDate(datetime.now().date())
            self.period_end_field = QDateEdit()
            self.period_end_field.setDate(datetime.now().date())
            self.kwh_field = QLineEdit()
            self.load_meters()
            layout.addRow("Счётчик:", self.meter_combo)
            layout.addRow("Период с:", self.period_start_field)
            layout.addRow("Период по:", self.period_end_field)
            layout.addRow("кВт·ч:", self.kwh_field)

        button_layout = QHBoxLayout()
        save_btn = QPushButton("Добавить")
        save_btn.clicked.connect(self.add_entity)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(save_btn)
        button_layout.addWidget(cancel_btn)
        layout.addRow(button_layout)
        self.setLayout(layout)

    def load_dropdowns(self):
        try:
            regions = requests.get(f"{API_BASE_URL}/regions", headers=HEADERS).json()
            tariffs = requests.get(f"{API_BASE_URL}/tariffs", headers=HEADERS).json()
            users = requests.get(f"{API_BASE_URL}/users", headers=HEADERS).json()
            if not isinstance(regions, list) or not isinstance(tariffs, list) or not isinstance(users, list):
                raise ValueError("Invalid data format from API")
            self.region_combo.clear()
            for r in regions:
                self.region_combo.addItem(r["name"], r["id"])
            self.tariff_combo.clear()
            for t in tariffs:
                self.tariff_combo.addItem(t["name"], t["id"])
            self.user_combo.clear()
            for u in users:
                self.user_combo.addItem(f"{u['login']} ({u['role']})", u["id"])
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные:\n{str(e)}")

    def load_buildings(self):
        try:
            if self.main_window.current_user_role == "tenant":
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(buildings, list):
                    raise ValueError("Invalid data format")
                user_id = self.main_window.current_user_id
                buildings = [b for b in buildings if b.get("user_id") == user_id]
            else:
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(buildings, list):
                    raise ValueError("Invalid data format")

            self.building_combo.clear()
            for b in buildings:
                self.building_combo.addItem(b["name"], b["id"])
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить объекты:\n{str(e)}")

    def load_meters(self):
        try:
            if self.main_window.current_user_role == "tenant":
                meters = requests.get(f"{API_BASE_URL}/meters", headers=HEADERS).json()
                buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
                if not isinstance(meters, list) or not isinstance(buildings, list):
                    raise ValueError("Invalid data format")
                user_id = self.main_window.current_user_id
                building_ids = [b["id"] for b in buildings if b.get("user_id") == user_id]
                meters = [m for m in meters if m.get("building_id") in building_ids]
            else:
                meters = requests.get(f"{API_BASE_URL}/meters", headers=HEADERS).json()
                if not isinstance(meters, list):
                    raise ValueError("Invalid data format")

            self.meter_combo.clear()
            for m in meters:
                self.meter_combo.addItem(m["serial_number"], m["id"])
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить счётчики:\n{str(e)}")

    def add_entity(self):
        data = {}
        if self.entity_type == "region":
            name = self.name_field.text().strip()
            tz = self.timezone_field.text().strip()
            if not name or not tz:
                QMessageBox.critical(self, "Ошибка", "Заполните все поля!")
                return
            data = {"name": name, "timezone": tz}
        elif self.entity_type == "tariff":
            name = self.name_field.text().strip()
            try:
                rate = float(self.rate_field.text().strip())
            except ValueError:
                QMessageBox.critical(self, "Ошибка", "Цена за кВт·ч должна быть числом!")
                return
            data = {
                "name": name,
                "rate_per_kwh": rate,
                "valid_from": self.valid_from_field.date().toString("yyyy-MM-dd"),
                "valid_to": self.valid_to_field.date().toString("yyyy-MM-dd") if self.valid_to_field.date().isValid() else None
            }
        elif self.entity_type == "user":
            login = self.login_field.text().strip()
            role = self.role_combo.currentText()
            if not login:
                QMessageBox.critical(self, "Ошибка", "Введите логин!")
                return
            data = {"login": login, "password_hash": "123",
                    "role_id": 1 if role == "tenant" else (2 if role == "accountant" else 3)}
        elif self.entity_type == "building":
            name = self.name_field.text().strip()
            addr = self.address_field.text().strip()
            region_id = self.region_combo.currentData()
            tariff_id = self.tariff_combo.currentData()
            user_id = self.user_combo.currentData()
            if not name or not addr or not region_id or not tariff_id or not user_id:
                QMessageBox.critical(self, "Ошибка", "Заполните все поля!")
                return
            data = {
                "name": name,
                "address": addr,
                "type": self.type_combo.currentText(),
                "region_id": region_id,
                "tariff_id": tariff_id,
                "user_id": user_id
            }
        elif self.entity_type == "meter":
            serial = self.serial_field.text().strip()
            date = self.installation_date_field.date().toString("yyyy-MM-dd")
            building_id = self.building_combo.currentData()
            if not serial or not date or not building_id:
                QMessageBox.critical(self, "Ошибка", "Заполните все поля!")
                return
            data = {
                "serial_number": serial,
                "installation_date": date,
                "building_id": building_id
            }
        elif self.entity_type == "consumption":
            meter_id = self.meter_combo.currentData()
            start = self.period_start_field.date().toString("yyyy-MM-dd")
            end = self.period_end_field.date().toString("yyyy-MM-dd")
            try:
                kwh = float(self.kwh_field.text().strip())
            except ValueError:
                QMessageBox.critical(self, "Ошибка", "Потребление должно быть числом!")
                return
            if not meter_id or not start or not end or not kwh:
                QMessageBox.critical(self, "Ошибка", "Заполните все поля!")
                return
            data = {
                "meter_id": meter_id,
                "period_start": start,
                "period_end": end,
                "consumption_kwh": kwh
            }

        try:
            if self.entity_type == "consumption":
                url = f"{API_BASE_URL}/consumption"
            else:
                url = f"{API_BASE_URL}/{self.entity_type}s"

            response = requests.post(url, json=data, headers=HEADERS, timeout=10)
            if response.status_code == 201:
                QMessageBox.information(self, "Успех", "Запись успешно добавлена!")
                self.accept()
                self.parent.refresh_data(self.entity_type)
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить запись:\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка подключения:\n{str(e)}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Система учёта электроэнергии")
        self.setGeometry(100, 100, 1200, 800)
        self.current_user_id = None
        self.current_user_role = None

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        header_layout = QHBoxLayout()
        self.title_label = QLabel("Система учёта электроэнергии")
        self.title_label.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        header_layout.addWidget(self.title_label)
        self.user_info_label = QLabel("Не авторизован")
        header_layout.addWidget(self.user_info_label)
        logout_btn = QPushButton("Выйти")
        logout_btn.clicked.connect(self.logout)
        header_layout.addWidget(logout_btn)
        main_layout.addLayout(header_layout)

        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)
        self.create_tabs()
        self.show_login_dialog()

    def show_login_dialog(self):
        dialog = LoginDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.user_info_label.setText(f"Вы вошли как: {self.current_user_role} (ID: {self.current_user_id})")
            self.update_ui_for_role(self.current_user_role)
            self.refresh_all_data()
        else:
            sys.exit()

    def update_ui_for_role(self, role):
        for i in range(self.tab_widget.count()):
            tab = self.tab_widget.widget(i)
            tab_name = self.tab_widget.tabText(i)
            if role == "admin":
                self.tab_widget.setTabVisible(i, True)
                self.tab_widget.setTabEnabled(i, True)
            elif role == "accountant":
                if tab_name == "Пользователи":
                    self.tab_widget.setTabVisible(i, False)
                else:
                    self.tab_widget.setTabVisible(i, True)
                    self.tab_widget.setTabEnabled(i, True)
            elif role == "tenant":
                visible = tab_name in ["Объекты учёта", "Счётчики", "Показания потребления"]
                self.tab_widget.setTabVisible(i, visible)
                self.tab_widget.setTabEnabled(i, visible)
            if hasattr(tab, "add_button"):
                tab.add_button.setVisible(role == "admin")

            if hasattr(tab, "delete_button"):
                tab.delete_button.setVisible(role == "admin")
            if tab_name == "Отчёты":
                visible = role == "accountant"
                self.tab_widget.setTabVisible(i, visible)

                if hasattr(tab, "generate_report_button"):
                    tab.generate_report_button.setVisible(visible)

    def create_tabs(self):
        self.tab_widget.addTab(self.create_entity_tab("region"), "Регионы")
        self.tab_widget.addTab(self.create_entity_tab("tariff"), "Тарифы")
        self.tab_widget.addTab(self.create_entity_tab("user"), "Пользователи")
        self.tab_widget.addTab(self.create_entity_tab("building"), "Объекты учёта")
        self.tab_widget.addTab(self.create_entity_tab("meter"), "Счётчики")
        self.tab_widget.addTab(self.create_entity_tab("consumption"), "Показания потребления")
        self.tab_widget.addTab(self.create_reports_tab(), "Отчёты")

    def create_entity_tab(self, entity_type):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        readable_names = {
            "region": "регион",
            "tariff": "тариф",
            "user": "пользователя",
            "building": "объект учёта",
            "meter": "счётчик",
            "consumption": "показание потребления"
        }

        button_layout = QHBoxLayout()

        add_button = QPushButton(f"Добавить {readable_names[entity_type]}")
        add_button.clicked.connect(lambda: self.add_entity(entity_type))
        button_layout.addWidget(add_button)

        delete_button = QPushButton("Удалить выбранное")
        delete_button.clicked.connect(
            lambda: self.delete_selected_entity(entity_type)
        )
        button_layout.addWidget(delete_button)

        # сохраним ссылки на кнопки
        widget.add_button = add_button
        widget.delete_button = delete_button

        layout.addLayout(button_layout)

        columns_map = {
            "region": ["ID", "Название", "Часовой пояс"],
            "tariff": ["ID", "Название", "Цена за кВт·ч", "С", "По"],
            "user": ["ID", "Логин", "Роль"],
            "building": ["ID", "Название", "Адрес", "Тип", "Регион", "Тариф", "Владелец"],
            "meter": ["ID", "Серийный номер", "Дата установки", "Объект"],
            "consumption": ["ID", "Счётчик", "Период с", "Период по", "кВт·ч", "Оценка (руб)"]
        }

        table = EntityTableWidget(widget, entity_type, columns_map[entity_type], self)
        layout.addWidget(table)
        setattr(widget, 'table', table)

        # Контекстное меню для удаления (только admin)
        if self.current_user_role == "admin":
            table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            table.customContextMenuRequested.connect(
                lambda pos, t=table, et=entity_type: self.show_context_menu(pos, t, et)
            )

        refresh_btn = QPushButton("Обновить")
        refresh_btn.clicked.connect(lambda: self.refresh_data(entity_type))
        layout.addWidget(refresh_btn)
        return widget

    def delete_selected_entity(self, entity_type):
        if self.current_user_role != "admin":
            QMessageBox.critical(self, "Ошибка", "Удалять могут только администраторы!")
            return

        # Найти нужную таблицу
        table = None
        for i in range(self.tab_widget.count()):
            tab = self.tab_widget.widget(i)
            if hasattr(tab, "table") and tab.table.entity_type == entity_type:
                table = tab.table
                break

        if not table:
            return

        selected_rows = table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Ошибка", "Выберите запись для удаления")
            return

        row = selected_rows[0].row()
        item = table.item(row, 0)
        if not item:
            return

        entity_id = int(item.text())
        self.delete_entity(entity_type, entity_id)
    def show_context_menu(self, pos, table, entity_type):
        row = table.rowAt(pos.y())
        if row < 0:
            return
        item = table.item(row, 0)
        if not item:
            return
        entity_id = int(item.text())

        context_menu = QMenu(self)
        delete_action = QAction("Удалить", self)
        delete_action.triggered.connect(lambda: self.delete_entity(entity_type, entity_id))
        context_menu.addAction(delete_action)
        context_menu.exec(table.viewport().mapToGlobal(pos))

    def delete_entity(self, entity_type, entity_id):
        if self.current_user_role != "admin":
            QMessageBox.critical(self, "Ошибка", "Удалять могут только администраторы!")
            return

        reply = QMessageBox.question(
            self,
            "Подтверждение",
            f"Вы уверены, что хотите удалить эту запись?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.No:
            return

        try:
            if entity_type == "consumption":
                url = f"{API_BASE_URL}/consumption/{entity_id}"
            else:
                url = f"{API_BASE_URL}/{entity_type}s/{entity_id}"

            response = requests.delete(url, headers=HEADERS, timeout=10)
            if response.status_code in (200, 204):
                QMessageBox.information(self, "Успех", "Запись успешно удалена!")
                self.refresh_data(entity_type)
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить запись:\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка подключения:\n{str(e)}")

    def refresh_data(self, entity_type):
        try:
            url = f"{API_BASE_URL}/{entity_type}s" if entity_type != "consumption" else f"{API_BASE_URL}/consumption"
            response = requests.get(url, headers=HEADERS, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if not isinstance(data, list):
                    error_msg = data.get("error", "Некорректный ответ сервера") if isinstance(data, dict) else str(data)
                    QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные:\n{error_msg}")
                    return
                for i in range(self.tab_widget.count()):
                    tab = self.tab_widget.widget(i)
                    tab_name = self.tab_widget.tabText(i)
                    enames = [("Регионы","region"), ("Тарифы","tariff"), ("Пользователи","user"),
                              ("Объекты учёта","building"), ("Счётчики","meter"), ("Показания потребления","consumption")]
                    if any(tab_name == name and etype == entity_type for name, etype in enames):
                        table = getattr(tab, 'table', None)
                        if table:
                            table.populate_table(data)
                        break
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные:\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось подключиться к серверу:\n{str(e)}")

    def refresh_all_data(self):
        role = self.current_user_role

        if role == "admin":
            entities = ["region", "tariff", "user", "building", "meter", "consumption"]

        elif role == "accountant":
            entities = ["tariff", "building", "meter", "consumption"]

        elif role == "tenant":
            entities = ["building", "meter", "consumption"]

        else:
            entities = []

        for et in entities:
            self.refresh_data(et)

    def add_entity(self, entity_type):
        # Проверка прав на создание
        if self.current_user_role == "tenant":
            return  # Арендатор ничего не может создавать
        if self.current_user_role == "accountant" and entity_type not in ["tariff", "consumption"]:
            return

        dialog = AddEntityDialog(self, entity_type, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            pass

    def edit_entity(self, entity_type, entity_id):
        try:
            if entity_type == "consumption":
                url = f"{API_BASE_URL}/consumption/{entity_id}"
            else:
                url = f"{API_BASE_URL}/{entity_type}s/{entity_id}"

            response = requests.get(url, headers=HEADERS, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if not isinstance(data, dict):
                    raise Exception("Сервер вернул некорректные данные")
                dialog = EditEntityDialog(self, entity_type, data, self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    pass
            else:
                error_msg = response.json().get("error", "Неизвестная ошибка")
                if response.status_code == 403:
                    error_msg = "Недостаточно прав для редактирования этой записи"
                QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить данные:\n{error_msg}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить данные для редактирования:\n{str(e)}")

    def logout(self):
        self.current_user_id = None
        self.current_user_role = None
        HEADERS.pop('X-User-ID', None)
        self.user_info_label.setText("Не авторизован")
        self.show_login_dialog()

    def generate_report(self):
        if self.current_user_role != "accountant":
            QMessageBox.critical(self, "Ошибка", "Доступ запрещён")
            return

        try:
            consumptions = requests.get(f"{API_BASE_URL}/consumption", headers=HEADERS).json()
            meters = requests.get(f"{API_BASE_URL}/meters", headers=HEADERS).json()
            buildings = requests.get(f"{API_BASE_URL}/buildings", headers=HEADERS).json()
            tariffs = requests.get(f"{API_BASE_URL}/tariffs", headers=HEADERS).json()

            meter_map = {m["id"]: m for m in meters}
            building_map = {b["id"]: b for b in buildings}
            tariff_map = {t["id"]: t for t in tariffs}

            report = {}
            total_cost = 0.0

            for c in consumptions:
                meter = meter_map.get(c["meter_id"])
                building = building_map.get(meter["building_id"])
                tariff = tariff_map.get(building["tariff_id"])

                kwh = c["consumption_kwh"]
                cost = kwh * tariff["rate_per_kwh"]
                name = building["name"]

                report.setdefault(name, {"kwh": 0.0, "cost": 0.0})
                report[name]["kwh"] += kwh
                report[name]["cost"] += cost
                total_cost += cost

            # Диалог сохранения
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Сохранить отчёт",
                "energy_report.docx",
                "Word files (*.docx)"
            )

            if not file_path:
                return

            # Создание Word-документа
            doc = Document()
            doc.add_heading("Отчёт по учёту электроэнергии", level=1)
            doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")

            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Объект"
            hdr_cells[1].text = "Потребление (кВт·ч)"
            hdr_cells[2].text = "Стоимость (руб.)"

            for name, data in report.items():
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = f"{data['kwh']:.2f}"
                row_cells[2].text = f"{data['cost']:.2f}"

            doc.add_paragraph(f"\nИтого: {total_cost:.2f} руб.")

            doc.save(file_path)

            QMessageBox.information(
                self,
                "Готово",
                f"Отчёт успешно сохранён:\n{file_path}"
            )

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка формирования отчёта:\n{str(e)}")

    def create_reports_tab(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        info = QLabel(
            "Формирование произвольного отчёта по учёту электроэнергии\n"
            "на основе данных о потреблении, тарифах и объектах."
        )
        layout.addWidget(info)

        generate_btn = QPushButton("Сформировать отчёт")
        generate_btn.clicked.connect(self.generate_report)
        layout.addWidget(generate_btn)

        # сохраняем ссылку для управления видимостью
        widget.generate_report_button = generate_btn

        layout.addStretch()
        return widget


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())